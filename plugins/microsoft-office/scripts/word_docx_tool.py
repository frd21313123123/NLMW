#!/usr/bin/env python3
"""Inspect and lightly edit Microsoft Word .docx files."""

from __future__ import annotations

import argparse
import json
import re
import sys
import zipfile
from collections import Counter
from pathlib import Path
from typing import Iterable


def require_docx():
    try:
        from docx import Document  # type: ignore
    except ImportError as exc:
        raise SystemExit(
            "Missing dependency: python-docx. Install it or load the bundled workspace dependencies."
        ) from exc
    return Document


def json_default(value):
    if hasattr(value, "isoformat"):
        return value.isoformat()
    return str(value)


def open_document(path: Path):
    Document = require_docx()
    if not path.exists():
        raise SystemExit(f"File not found: {path}")
    if path.suffix.lower() != ".docx":
        raise SystemExit("This helper supports .docx files. Convert legacy .doc files first.")
    return Document(str(path))


def iter_paragraphs(document) -> Iterable[tuple[str, object]]:
    for index, paragraph in enumerate(document.paragraphs):
        yield f"paragraph[{index}]", paragraph
    for table_index, table in enumerate(document.tables):
        for row_index, row in enumerate(table.rows):
            for cell_index, cell in enumerate(row.cells):
                for paragraph_index, paragraph in enumerate(cell.paragraphs):
                    location = (
                        f"table[{table_index}].row[{row_index}]."
                        f"cell[{cell_index}].paragraph[{paragraph_index}]"
                    )
                    yield location, paragraph


def paragraph_text(paragraph) -> str:
    return paragraph.text.strip()


def table_to_rows(table) -> list[list[str]]:
    return [[cell.text.strip() for cell in row.cells] for row in table.rows]


def zip_features(path: Path) -> dict[str, object]:
    features = {
        "media_count": 0,
        "has_comments": False,
        "has_footnotes": False,
        "has_endnotes": False,
        "has_numbering": False,
    }
    with zipfile.ZipFile(path) as archive:
        names = set(archive.namelist())
        features["media_count"] = sum(1 for name in names if name.startswith("word/media/"))
        features["has_comments"] = "word/comments.xml" in names
        features["has_footnotes"] = "word/footnotes.xml" in names
        features["has_endnotes"] = "word/endnotes.xml" in names
        features["has_numbering"] = "word/numbering.xml" in names
    return features


def inspect_docx(args) -> int:
    path = Path(args.docx)
    document = open_document(path)
    props = document.core_properties
    styles = Counter()
    non_empty_paragraphs = 0
    headings = []

    for location, paragraph in iter_paragraphs(document):
        style_name = getattr(paragraph.style, "name", "") or ""
        if style_name:
            styles[style_name] += 1
        text = paragraph_text(paragraph)
        if text:
            non_empty_paragraphs += 1
        if text and style_name.lower().startswith("heading"):
            headings.append({"location": location, "style": style_name, "text": text})

    tables = []
    for index, table in enumerate(document.tables):
        row_count = len(table.rows)
        column_count = len(table.columns)
        preview = table_to_rows(table)[:3]
        tables.append(
            {
                "index": index,
                "rows": row_count,
                "columns": column_count,
                "preview": preview,
            }
        )

    result = {
        "file": str(path),
        "paragraphs": len(document.paragraphs),
        "non_empty_paragraphs": non_empty_paragraphs,
        "tables": tables,
        "sections": len(document.sections),
        "headings": headings,
        "styles_used": dict(styles.most_common()),
        "core_properties": {
            "title": props.title,
            "subject": props.subject,
            "author": props.author,
            "keywords": props.keywords,
            "comments": props.comments,
            "created": props.created,
            "modified": props.modified,
            "last_modified_by": props.last_modified_by,
        },
        "package_features": zip_features(path),
    }
    print(json.dumps(result, ensure_ascii=False, indent=2, default=json_default))
    return 0


def extract_docx(args) -> int:
    path = Path(args.docx)
    document = open_document(path)
    items = []

    for location, paragraph in iter_paragraphs(document):
        text = paragraph_text(paragraph)
        if not text:
            continue
        style_name = getattr(paragraph.style, "name", "") or ""
        items.append({"type": "paragraph", "location": location, "style": style_name, "text": text})

    for index, table in enumerate(document.tables):
        items.append({"type": "table", "location": f"table[{index}]", "rows": table_to_rows(table)})

    if args.format == "json":
        output = json.dumps(items, ensure_ascii=False, indent=2)
    elif args.format == "markdown":
        lines = []
        for item in items:
            if item["type"] == "paragraph":
                style = str(item.get("style", ""))
                text = str(item["text"])
                match = re.match(r"Heading\s+(\d+)", style, flags=re.I)
                if match:
                    level = min(max(int(match.group(1)), 1), 6)
                    lines.append(f"{'#' * level} {text}")
                else:
                    lines.append(text)
            else:
                rows = item["rows"]
                if rows:
                    width = max(len(row) for row in rows)
                    padded = [row + [""] * (width - len(row)) for row in rows]
                    lines.append("")
                    lines.append("| " + " | ".join(padded[0]) + " |")
                    lines.append("| " + " | ".join(["---"] * width) + " |")
                    for row in padded[1:]:
                        lines.append("| " + " | ".join(row) + " |")
                    lines.append("")
        output = "\n\n".join(lines)
    else:
        output = "\n\n".join(item["text"] for item in items if item["type"] == "paragraph")

    if args.limit_chars and len(output) > args.limit_chars:
        output = output[: args.limit_chars] + "\n...[truncated]"
    print(output)
    return 0


def find_docx(args) -> int:
    pattern = re.compile(re.escape(args.pattern), 0 if args.case_sensitive else re.I)
    document = open_document(Path(args.docx))
    matches = []
    for location, paragraph in iter_paragraphs(document):
        text = paragraph.text
        if pattern.search(text):
            matches.append({"location": location, "text": text.strip()})
    print(json.dumps(matches, ensure_ascii=False, indent=2))
    return 0 if matches else 1


def stats_docx(args) -> int:
    document = open_document(Path(args.docx))
    text_parts = [paragraph.text for _, paragraph in iter_paragraphs(document)]
    text = "\n".join(text_parts)
    words = re.findall(r"\b\w+\b", text, flags=re.UNICODE)
    result = {
        "characters": len(text),
        "characters_no_spaces": len(re.sub(r"\s+", "", text)),
        "words": len(words),
        "paragraphs_with_tables": len(list(iter_paragraphs(document))),
        "body_paragraphs": len(document.paragraphs),
        "tables": len(document.tables),
        "sections": len(document.sections),
    }
    print(json.dumps(result, ensure_ascii=False, indent=2))
    return 0


def replace_in_text(text: str, find: str, replacement: str, case_sensitive: bool, limit: int | None):
    flags = 0 if case_sensitive else re.I
    pattern = re.compile(re.escape(find), flags)
    return pattern.subn(replacement, text, count=0 if limit is None else limit)


def replace_docx(args) -> int:
    input_path = Path(args.input_docx)
    output_path = Path(args.output_docx)
    if input_path.resolve() == output_path.resolve() and not args.allow_overwrite:
        raise SystemExit("Refusing to overwrite the input file without --allow-overwrite.")

    document = open_document(input_path)
    remaining = args.max_count
    count = 0

    for _, paragraph in iter_paragraphs(document):
        for run in paragraph.runs:
            if remaining is not None and remaining <= 0:
                break
            new_text, changed = replace_in_text(
                run.text,
                args.find,
                args.replace,
                args.case_sensitive,
                remaining,
            )
            if changed:
                run.text = new_text
                count += changed
                if remaining is not None:
                    remaining -= changed

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))
    print(json.dumps({"replacements": count, "output": str(output_path)}, indent=2))
    return 0


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description=__doc__)
    subparsers = parser.add_subparsers(dest="command", required=True)

    inspect_parser = subparsers.add_parser("inspect", help="Inspect document structure.")
    inspect_parser.add_argument("docx")
    inspect_parser.set_defaults(func=inspect_docx)

    extract_parser = subparsers.add_parser("extract", help="Extract document text and tables.")
    extract_parser.add_argument("docx")
    extract_parser.add_argument("--format", choices=["text", "markdown", "json"], default="text")
    extract_parser.add_argument("--limit-chars", type=int, default=0)
    extract_parser.set_defaults(func=extract_docx)

    find_parser = subparsers.add_parser("find", help="Find exact text in paragraphs and table cells.")
    find_parser.add_argument("docx")
    find_parser.add_argument("pattern")
    find_parser.add_argument("--case-sensitive", action="store_true")
    find_parser.set_defaults(func=find_docx)

    stats_parser = subparsers.add_parser("stats", help="Compute basic document statistics.")
    stats_parser.add_argument("docx")
    stats_parser.set_defaults(func=stats_docx)

    replace_parser = subparsers.add_parser("replace", help="Run-level find/replace that preserves local formatting when matches are inside individual runs.")
    replace_parser.add_argument("input_docx")
    replace_parser.add_argument("output_docx")
    replace_parser.add_argument("--find", required=True)
    replace_parser.add_argument("--replace", required=True)
    replace_parser.add_argument("--case-sensitive", action="store_true")
    replace_parser.add_argument("--max-count", type=int)
    replace_parser.add_argument("--allow-overwrite", action="store_true")
    replace_parser.set_defaults(func=replace_docx)

    return parser


def main(argv: list[str] | None = None) -> int:
    parser = build_parser()
    args = parser.parse_args(argv)
    return args.func(args)


if __name__ == "__main__":
    sys.exit(main())
